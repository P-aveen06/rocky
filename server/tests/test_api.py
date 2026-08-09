from __future__ import annotations

import asyncio
import inspect
import io
import json
import socket
import sqlite3
import ssl
import threading
import time
from concurrent.futures import ThreadPoolExecutor
from datetime import UTC, datetime, timedelta
from functools import lru_cache
from pathlib import Path

import jwt
import pytest
from alembic import command
from alembic.config import Config
from cryptography.hazmat.primitives import serialization
from cryptography.hazmat.primitives.asymmetric import rsa
from docx import Document
from fastapi.testclient import TestClient
from pydantic import ValidationError
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject
from sqlalchemy.ext.asyncio import AsyncSession

import api.database as database
import api.routes.realtime as realtime_routes
from api.config import REPO_ROOT, SERVER_ROOT, Settings, get_settings
from api.database import (
    database_connect_args,
    normalized_database_url,
)
from api.main import create_app
from api.models import InterviewTurn
from api.services.realtime import RealtimeClientSecret
from api.services.transcription import FinalTranscription, TranscriptionServiceError
from domain.evaluation import (
    CompetencyEvaluation,
    EvaluationReport,
    EvidenceCitation,
)

_CLERK_ISSUER = "https://clerk.example.test"
_CLERK_AUTHORIZED_PARTY = "https://app.example.test"


@lru_cache(maxsize=1)
def _clerk_keys() -> tuple[str, str]:
    """A throwaway RSA pair standing in for a Clerk instance's signing key.

    Handing the public half to the app as ``clerk_jwt_key`` keeps verification
    networkless, so the suite never reaches for a JWKS endpoint.
    """

    private_key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
    private_pem = private_key.private_bytes(
        serialization.Encoding.PEM,
        serialization.PrivateFormat.PKCS8,
        serialization.NoEncryption(),
    ).decode()
    public_pem = (
        private_key.public_key()
        .public_bytes(
            serialization.Encoding.PEM,
            serialization.PublicFormat.SubjectPublicKeyInfo,
        )
        .decode()
    )
    return private_pem, public_pem


def _clerk_settings(**overrides: object) -> Settings:
    _, public_pem = _clerk_keys()
    values: dict[str, object] = {
        "_env_file": None,
        "app_env": "test",
        "auth_mode": "clerk",
        "clerk_secret_key": "sk_test_example",
        "clerk_publishable_key": "pk_test_example",
        "clerk_frontend_api_url": _CLERK_ISSUER,
        "clerk_jwt_key": public_pem,
        "clerk_authorized_parties": _CLERK_AUTHORIZED_PARTY,
    }
    values.update(overrides)
    return Settings(**values)


def _clerk_claims(subject: str, email: str) -> dict[str, object]:
    now = int(time.time())
    return {
        "sub": subject,
        "email": email,
        "name": email.split("@", 1)[0],
        "iss": _CLERK_ISSUER,
        "azp": _CLERK_AUTHORIZED_PARTY,
        "iat": now,
        "nbf": now,
        "exp": now + 3600,
    }


def _principal_headers(subject: str, email: str) -> dict[str, str]:
    private_pem, _ = _clerk_keys()
    token = jwt.encode(_clerk_claims(subject, email), private_pem, algorithm="RS256")
    return {"Authorization": f"Bearer {token}"}


@pytest.mark.real_env_file
def test_settings_load_optional_local_env_after_base_env(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    base_env = tmp_path / ".env"
    local_env = tmp_path / ".env.local"
    base_env.write_text(
        "DATABASE_URL=postgresql+asyncpg://user:password@database/app\n"
        "AUTO_CREATE_SCHEMA=false\n",
        encoding="utf-8",
    )
    local_env.write_text(
        "DATABASE_URL=sqlite+aiosqlite:///./data/interview_coach.db\n"
        "AUTO_CREATE_SCHEMA=true\n",
        encoding="utf-8",
    )
    monkeypatch.delenv("DATABASE_URL", raising=False)
    monkeypatch.delenv("AUTO_CREATE_SCHEMA", raising=False)

    settings = Settings(_env_file=(base_env, local_env))

    assert settings.database_url == "sqlite+aiosqlite:///./data/interview_coach.db"
    assert settings.auto_create_schema is True
    assert Settings.model_config["env_file"] == (
        REPO_ROOT / ".env",
        REPO_ROOT / ".env.local",
    )


def test_settings_process_environment_overrides_dotenv_files(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    base_env = tmp_path / ".env"
    local_env = tmp_path / ".env.local"
    base_env.write_text(
        "DATABASE_URL=postgresql+asyncpg://base:password@database/app\n"
        "AUTO_CREATE_SCHEMA=false\n",
        encoding="utf-8",
    )
    local_env.write_text(
        "DATABASE_URL=sqlite+aiosqlite:///./data/interview_coach.db\n"
        "AUTO_CREATE_SCHEMA=true\n",
        encoding="utf-8",
    )
    monkeypatch.setenv(
        "DATABASE_URL", "postgresql+asyncpg://process:password@database/app"
    )
    monkeypatch.setenv("AUTO_CREATE_SCHEMA", "false")

    settings = Settings(_env_file=(base_env, local_env))

    assert settings.database_url == "postgresql+asyncpg://process:password@database/app"
    assert settings.auto_create_schema is False


def test_settings_reject_local_auth_and_sqlite_in_staging(tmp_path: Path) -> None:
    with pytest.raises(ValidationError):
        Settings(
            app_env="staging",
            auth_mode="local",
            database_url=f"sqlite+aiosqlite:///{tmp_path / 'staging.db'}",
            auto_create_schema=False,
        )


def test_staging_accepts_managed_auth_with_postgresql() -> None:
    settings = _clerk_settings(
        app_env="staging",
        database_url="postgresql+asyncpg://user:password@database/app",
        auto_create_schema=False,
    )

    assert settings.auth_mode == "clerk"


def test_clerk_auth_mode_requires_clerk_credentials() -> None:
    with pytest.raises(ValidationError):
        Settings(_env_file=None, app_env="test", auth_mode="clerk")


def test_neon_url_is_safe_for_asyncpg() -> None:
    url = normalized_database_url(
        "postgresql://candidate:secret@ep-example.us-east-2.aws.neon.tech/"
        "interview_coach?sslmode=require&channel_binding=require"
    )

    assert url == (
        "postgresql+asyncpg://candidate:secret@"
        "ep-example.us-east-2.aws.neon.tech/interview_coach"
    )
    connect_args = database_connect_args(url, timeout_seconds=7.5)
    ssl_context = connect_args["ssl"]
    assert isinstance(ssl_context, ssl.SSLContext)
    assert ssl_context.check_hostname is True
    assert ssl_context.verify_mode == ssl.CERT_REQUIRED
    assert connect_args["timeout"] == 7.5


@pytest.mark.asyncio
async def test_neon_connection_prefers_ipv4_without_changing_tls_hostname() -> None:
    calls: list[tuple[str | None, int | None, dict[str, object]]] = []

    class FakeLoop:
        async def create_connection(
            self,
            protocol_factory,
            host: str | None = None,
            port: int | None = None,
            **kwargs: object,
        ) -> str:
            calls.append((host, port, kwargs))
            return "connected"

    hostname = "ep-example-pooler.us-east-2.aws.neon.tech"
    loop = FakeLoop()
    database.install_database_network_compatibility(
        f"postgresql+asyncpg://candidate:secret@{hostname}/interview_coach",
        loop=loop,
    )

    result = await loop.create_connection(object, hostname, 5432, ssl="verified")

    assert result == "connected"
    assert calls == [(hostname, 5432, {"ssl": "verified", "family": socket.AF_INET})]


@pytest.fixture
def settings(tmp_path: Path) -> Settings:
    return Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'test.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        enable_text_dev_mode=False,
    )


@pytest.fixture
def client(settings: Settings):
    with TestClient(create_app(settings)) as test_client:
        yield test_client


def _create_session(
    client: TestClient, *, headers: dict[str, str] | None = None
) -> dict[str, object]:
    response = client.post(
        "/api/interviews",
        json={"title": "Untitled practice session"},
        headers=headers,
    )
    assert response.status_code == 201
    return response.json()


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Alex Morgan")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, FastAPI, PostgreSQL")
    document.add_paragraph("Experience")
    document.add_paragraph(
        "Built payment APIs and reduced database query latency by 35%."
    )
    document.add_paragraph("Education")
    document.add_paragraph("BSc Computer Science")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _ready_session(
    client: TestClient, *, headers: dict[str, str] | None = None
) -> dict[str, object]:
    interview = _create_session(client, headers=headers)
    upload = client.post(
        "/api/uploads/resume",
        data={"interview_id": interview["id"]},
        files={
            "file": (
                "alex.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document",
            )
        },
        headers=headers,
    ).json()
    client.post(
        "/api/candidate-profiles/extract",
        json={"interview_id": interview["id"], "upload_id": upload["id"]},
        headers=headers,
    )
    target = client.post(
        "/api/job-targets",
        json={
            "interview_id": interview["id"],
            "title": "Backend Engineer",
            "seniority": "mid",
            "raw_description": (
                "Build reliable Python and FastAPI services with PostgreSQL, "
                "testing, observability, incident response, and team ownership."
            ),
        },
        headers=headers,
    ).json()
    client.post(
        "/api/scorecards/generate",
        json={"interview_id": interview["id"], "job_target_id": target["id"]},
        headers=headers,
    )
    return interview


def _start_ready_interview(
    client: TestClient,
    interview_id: object,
    *,
    input_mode: str,
    headers: dict[str, str] | None = None,
) -> None:
    secret = client.post(
        f"/api/interviews/{interview_id}/realtime-client-secret",
        json={
            "input_mode": input_mode,
            "duration_minutes": 15,
            "interview_type": "technical_behavioral",
        },
        headers=headers,
    )
    assert secret.status_code == 200
    connected = client.post(
        f"/api/interviews/{interview_id}/connection-state",
        json={"state": "connected"},
        headers=headers,
    )
    assert connected.status_code == 200


async def _fake_realtime_secret(**_kwargs: object) -> RealtimeClientSecret:
    return RealtimeClientSecret(
        value="ek_temporary",
        expires_at=2_000_000_000,
        calls_url="https://example.invalid/openai/v1/realtime/calls",
    )


def _dual_transcription_settings(database_path: Path) -> Settings:
    return Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{database_path}",
        auto_create_schema=True,
        web_dist_dir=database_path.parent / "missing-dist",
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="server-key",
        azure_openai_realtime_deployment="interviewer-deployment",
        azure_openai_realtime_transcription_model="live-stt-deployment",
        azure_openai_final_transcription_deployment="final-stt-deployment",
    )


def _pdf_bytes(*, text: str | None = None, encrypted: bool = False) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    if text:
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        font_reference = writer._add_object(font)
        page[NameObject("/Resources")] = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_reference})}
        )
        stream = DecodedStreamObject()
        safe_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream.set_data(f"BT /F1 12 Tf 72 720 Td ({safe_text}) Tj ET".encode())
        page[NameObject("/Contents")] = writer._add_object(stream)
    if encrypted:
        writer.encrypt("secret")
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def test_health_reports_database_readiness(client: TestClient) -> None:
    health = client.get("/api/health")
    assert health.json() == {"status": "ok"}
    assert health.headers["X-Content-Type-Options"] == "nosniff"
    assert health.headers["Referrer-Policy"] == "no-referrer"
    assert "camera=()" in health.headers["Permissions-Policy"]
    assert "frame-ancestors 'none'" in health.headers["Content-Security-Policy"]
    assert client.get("/api/health/ready").json() == {
        "status": "ready",
        "database": "ok",
    }


def test_user_creates_session_and_sees_it_after_refresh(client: TestClient) -> None:
    user = client.get("/api/auth/me")
    assert user.status_code == 200
    assert user.json()["email"] == "developer@local.test"

    created = client.post(
        "/api/interviews",
        json={"title": "Untitled practice session"},
    )
    assert created.status_code == 201
    assert created.json()["status"] == "DRAFT"

    refreshed = client.get("/api/interviews")
    assert refreshed.status_code == 200
    assert [item["id"] for item in refreshed.json()["items"]] == [created.json()["id"]]


def test_m2_resume_profile_jd_and_editable_scorecard_flow(client: TestClient) -> None:
    interview = _create_session(client)
    uploaded = client.post(
        "/api/uploads/resume",
        data={"interview_id": interview["id"]},
        files={
            "file": (
                "alex-morgan.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document",
            )
        },
    )
    assert uploaded.status_code == 201
    assert uploaded.json()["file_type"] == "docx"
    assert uploaded.json()["raw_deleted_at"]

    profile = client.post(
        "/api/candidate-profiles/extract",
        json={
            "interview_id": interview["id"],
            "upload_id": uploaded.json()["id"],
        },
    )
    assert profile.status_code == 201
    profile_body = profile.json()
    assert profile_body["headline"] == "Alex Morgan"
    assert profile_body["claims"]
    assert profile_body["claims"][0]["source"]["source_id"].startswith("resume:block:")

    corrected_claims = [
        {"id": claim["id"], "text": claim["text"]} for claim in profile_body["claims"]
    ]
    corrected_claims[0]["text"] = "Python, FastAPI, PostgreSQL, and Redis"
    corrected = client.patch(
        f"/api/candidate-profiles/{profile_body['id']}",
        json={"headline": "Backend engineer", "claims": corrected_claims},
    )
    assert corrected.status_code == 200
    assert corrected.json()["claims"][0]["edited"] is True
    assert (
        corrected.json()["claims"][0]["original_text"]
        == (profile_body["claims"][0]["text"])
    )

    protected = client.post(
        "/api/candidate-profiles/extract",
        json={
            "interview_id": interview["id"],
            "upload_id": uploaded.json()["id"],
            "replace_existing": True,
        },
    )
    assert protected.status_code == 409
    assert "saved corrections" in protected.json()["error"]["message"]

    extracted_again = client.post(
        "/api/candidate-profiles/extract",
        json={
            "interview_id": interview["id"],
            "upload_id": uploaded.json()["id"],
        },
    )
    assert extracted_again.json()["headline"] == "Backend engineer"
    assert extracted_again.json()["claims"][0]["text"].endswith("and Redis")

    description = (
        "Senior Backend Engineer\n"
        "Python and FastAPI are required for production API development.\n"
        "Strong PostgreSQL and database design experience is essential.\n"
        "AWS or another cloud platform is preferred.\n"
        "Own testing, observability, incidents, and cross-team delivery."
    )
    target = client.post(
        "/api/job-targets",
        json={
            "interview_id": interview["id"],
            "title": "Senior Backend Engineer",
            "seniority": "senior",
            "raw_description": description,
        },
    )
    assert target.status_code == 201
    assert target.json()["structured_requirements"]

    scorecard = client.post(
        "/api/scorecards/generate",
        json={
            "interview_id": interview["id"],
            "job_target_id": target.json()["id"],
        },
    )
    assert scorecard.status_code == 201
    scorecard_body = scorecard.json()
    assert sum(item["weight"] for item in scorecard_body["competencies"]) == 100
    assert all(item["source_references"] for item in scorecard_body["competencies"])
    assert (
        "Shapes system boundaries"
        in scorecard_body["competencies"][0]["seniority_expectation"]
    )

    edits = []
    for competency in scorecard_body["competencies"]:
        edits.append(
            {
                key: competency[key]
                for key in (
                    "id",
                    "name",
                    "description",
                    "weight",
                    "classification",
                    "seniority_expectation",
                    "evidence_to_collect",
                    "question_families",
                )
            }
        )
    edits[0]["name"] = "Backend API design"
    saved_scorecard = client.patch(
        f"/api/scorecards/{scorecard_body['id']}", json={"competencies": edits}
    )
    assert saved_scorecard.status_code == 200
    assert saved_scorecard.json()["version"] == 2
    assert saved_scorecard.json()["competencies"][0]["name"] == "Backend API design"

    setup = client.get(f"/api/interviews/{interview['id']}/setup")
    assert setup.status_code == 200
    assert setup.json()["profile"]["headline"] == "Backend engineer"
    assert setup.json()["scorecard"]["total_weight"] == 100

    refreshed_session = client.get(f"/api/interviews/{interview['id']}")
    assert refreshed_session.json()["status"] == "SCORECARD_READY"
    assert refreshed_session.json()["profile_id"] == profile_body["id"]


@pytest.mark.parametrize(
    ("filename", "media_type", "contents", "expected_status", "message"),
    [
        (
            "spoofed.pdf",
            "application/pdf",
            b"not really a PDF",
            415,
            "does not contain a PDF signature",
        ),
        (
            "resume.docx",
            "application/pdf",
            _docx_bytes(),
            415,
            "content type",
        ),
        (
            "corrupt.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            b"PKcorrupt",
            422,
            "corrupt",
        ),
        (
            "locked.pdf",
            "application/pdf",
            _pdf_bytes(text="Secret resume", encrypted=True),
            422,
            "Encrypted",
        ),
        (
            "scanned.pdf",
            "application/pdf",
            _pdf_bytes(),
            422,
            "No selectable text",
        ),
        (
            "oversized.pdf",
            "application/pdf",
            b"%PDF-" + b"x" * 5_000_000,
            413,
            "5 MB limit",
        ),
    ],
)
def test_resume_upload_rejects_unsafe_or_unsupported_files(
    client: TestClient,
    filename: str,
    media_type: str,
    contents: bytes,
    expected_status: int,
    message: str,
) -> None:
    interview = _create_session(client)
    response = client.post(
        "/api/uploads/resume",
        data={"interview_id": interview["id"]},
        files={"file": (filename, contents, media_type)},
    )

    assert response.status_code == expected_status
    assert message in response.json()["error"]["message"]


def test_text_pdf_upload_succeeds_with_page_source(client: TestClient) -> None:
    interview = _create_session(client)
    response = client.post(
        "/api/uploads/resume",
        data={"interview_id": interview["id"]},
        files={
            "file": (
                "resume.pdf",
                _pdf_bytes(text="Alex Morgan Backend Engineer"),
                "application/pdf",
            )
        },
    )
    assert response.status_code == 201
    profile = client.post(
        "/api/candidate-profiles/extract",
        json={
            "interview_id": interview["id"],
            "upload_id": response.json()["id"],
        },
    )
    assert profile.status_code == 201
    assert profile.json()["headline"] == "Alex Morgan Backend Engineer"


def test_scorecard_rejects_weights_that_do_not_total_one_hundred(
    client: TestClient,
) -> None:
    interview = _create_session(client)
    target = client.post(
        "/api/job-targets",
        json={
            "interview_id": interview["id"],
            "title": "Backend Engineer",
            "seniority": "mid",
            "raw_description": (
                "Backend engineer required to design APIs, use SQL databases, "
                "test services, debug incidents, and collaborate with a team."
            ),
        },
    ).json()
    scorecard = client.post(
        "/api/scorecards/generate",
        json={"interview_id": interview["id"], "job_target_id": target["id"]},
    ).json()
    edits = []
    for competency in scorecard["competencies"]:
        edit = {
            key: competency[key]
            for key in (
                "id",
                "name",
                "description",
                "weight",
                "classification",
                "seniority_expectation",
                "evidence_to_collect",
                "question_families",
            )
        }
        edits.append(edit)
    edits[0]["weight"] += 1

    response = client.patch(
        f"/api/scorecards/{scorecard['id']}", json={"competencies": edits}
    )
    assert response.status_code == 422


def test_text_dev_mode_is_rejected_when_server_flag_is_off(client: TestClient) -> None:
    interview = _create_session(client)
    response = client.post(
        f"/api/interviews/{interview['id']}/realtime-client-secret",
        json={
            "input_mode": "text_dev",
            "duration_minutes": 15,
            "interview_type": "technical_behavioral",
        },
    )

    assert response.status_code == 403
    assert "disabled by the server" in response.json()["error"]["message"]


def test_capabilities_expose_only_safe_dual_transcription_configuration(
    tmp_path: Path,
) -> None:
    capability_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'capabilities.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="server-key",
        azure_openai_realtime_deployment="interviewer-deployment",
        azure_openai_realtime_transcription_model="live-stt-deployment",
        azure_openai_final_transcription_deployment="final-stt-deployment",
    )

    with TestClient(create_app(capability_settings)) as capability_client:
        response = capability_client.get("/api/capabilities")

    assert response.status_code == 200
    body = response.json()
    assert body["live_transcription_configured"] is True
    assert body["final_transcription_configured"] is True
    serialized_body = json.dumps(body)
    assert "interviewer-deployment" not in serialized_body
    assert "live-stt-deployment" not in serialized_body
    assert "final-stt-deployment" not in serialized_body
    assert "example.services.ai.azure.com" not in serialized_body
    assert "server-key" not in serialized_body


def test_final_transcription_openapi_documents_bounded_raw_multipart(
    client: TestClient,
) -> None:
    signature = inspect.signature(realtime_routes.finalize_candidate_transcription)

    assert "file" not in signature.parameters
    assert all(
        "UploadFile" not in str(parameter.annotation)
        for parameter in signature.parameters.values()
    )
    response = client.get("/api/openapi.json")
    assert response.status_code == 200
    operation = response.json()["paths"][
        "/api/interviews/{interview_id}/turns/{client_turn_id}:transcribe"
    ]["post"]
    request_body = operation["requestBody"]
    assert request_body["required"] is True
    schema = request_body["content"]["multipart/form-data"]["schema"]
    assert schema["required"] == ["file"]
    assert schema["properties"] == {
        "file": {"type": "string", "format": "binary"},
        "started_at": {"type": "string", "format": "date-time"},
        "ended_at": {"type": "string", "format": "date-time"},
    }
    assert operation["responses"]["200"]["content"]["application/json"]["schema"] == {
        "$ref": "#/components/schemas/InterviewRuntimeResponse"
    }
    accept_operation = response.json()["paths"][
        "/api/interviews/{interview_id}/turns/{client_turn_id}:accept-live"
    ]["post"]
    assert accept_operation["responses"]["200"]["content"]["application/json"][
        "schema"
    ] == {"$ref": "#/components/schemas/InterviewRuntimeResponse"}


@pytest.mark.parametrize("client_turn_id", ["ab", "bad.id", "x" * 97])
def test_turn_path_ids_share_the_batch_identifier_constraint(
    client: TestClient, client_turn_id: str
) -> None:
    interview = _create_session(client)

    transcribe = client.post(
        f"/api/interviews/{interview['id']}/turns/{client_turn_id}:transcribe",
        files={"file": ("answer.webm", b"audio", "audio/webm")},
    )
    accept = client.post(
        f"/api/interviews/{interview['id']}/turns/{client_turn_id}:accept-live"
    )

    assert transcribe.status_code == accept.status_code == 422


def test_final_transcription_upgrades_live_turn_once_and_late_live_is_noop(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    database_path = tmp_path / "dual-convergence.db"
    transcription_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{database_path}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="server-key",
        azure_openai_realtime_deployment="interviewer-deployment",
        azure_openai_realtime_transcription_model="live-stt-deployment",
        azure_openai_final_transcription_deployment="final-stt-deployment",
    )
    provider_calls: list[dict[str, object]] = []

    async def fake_transcription(**kwargs: object) -> FinalTranscription:
        provider_calls.append(kwargs)
        return FinalTranscription(
            text="I designed an idempotent payment API.",
            deployment="final-stt-deployment",
            elapsed_ms=2_400,
            attempts=2,
        )

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    monkeypatch.setattr(
        "api.routes.realtime.transcribe_candidate_audio",
        fake_transcription,
    )
    with TestClient(create_app(transcription_settings)) as transcription_client:
        interview = _ready_session(transcription_client)
        _start_ready_interview(
            transcription_client, interview["id"], input_mode="voice"
        )
        live = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json={
                "items": [
                    {
                        "client_turn_id": "assistant-question",
                        "speaker": "assistant",
                        "transcript": "Describe an API you designed.",
                    },
                    {
                        "client_turn_id": "candidate-answer",
                        "speaker": "user",
                        "transcript": "I designed a payment service.",
                    },
                ]
            },
        )
        assert live.status_code == 200
        assistant_turn, live_turn = live.json()["turns"]
        assert assistant_turn["transcription_source"] == "assistant"
        assert assistant_turn["transcription_model"] is None
        assert assistant_turn["transcription_finalized_at"] is not None
        assert live_turn["transcription_source"] == "realtime_live"
        assert live_turn["transcription_model"] == "live-stt-deployment"
        assert live_turn["transcription_finalized_at"] is None

        finalized = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns/candidate-answer:transcribe",
            files={
                "file": (
                    "candidate.webm",
                    b"candidate-audio",
                    "audio/webm; codecs=opus",
                )
            },
        )
        repeated = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns/candidate-answer:transcribe",
            files={"file": ("retry.webm", b"different-audio", "audio/webm")},
        )

        assert finalized.status_code == repeated.status_code == 200
        final_turn = next(
            turn
            for turn in finalized.json()["turns"]
            if turn["client_turn_id"] == "candidate-answer"
        )
        repeated_turn = next(
            turn
            for turn in repeated.json()["turns"]
            if turn["client_turn_id"] == "candidate-answer"
        )
        assert final_turn["id"] == live_turn["id"]
        assert final_turn["sequence"] == live_turn["sequence"]
        assert final_turn["transcript"] == "I designed an idempotent payment API."
        assert final_turn["transcription_source"] == "final_model"
        assert final_turn["transcription_model"] == "final-stt-deployment"
        assert final_turn["transcription_finalized_at"] is not None
        assert repeated_turn == final_turn
        assert len(provider_calls) == 1
        assert provider_calls[0]["audio"] == b"candidate-audio"
        assert provider_calls[0]["media_type"] == "audio/webm"

        late_live = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json={
                "items": [
                    {
                        "client_turn_id": "candidate-answer",
                        "speaker": "user",
                        "transcript": "A late and less accurate live transcript.",
                    }
                ]
            },
        )
        assert late_live.status_code == 200
        saved_turn = next(
            turn
            for turn in late_live.json()["turns"]
            if turn["client_turn_id"] == "candidate-answer"
        )
        assert saved_turn == final_turn

        pre_boundary_live = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json={
                "items": [
                    {
                        "client_turn_id": "pre-boundary-live",
                        "speaker": "user",
                        "transcript": "This live turn existed before completion.",
                    }
                ]
            },
        )
        assert pre_boundary_live.status_code == 200
        transcription_client.post(f"/api/interviews/{interview['id']}/complete")
        finalized_after_end = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns/pre-boundary-live:transcribe",
            files={"file": ("saved.webm", b"saved-audio", "audio/webm")},
        )
        after_end = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns/candidate-answer:transcribe",
            files={"file": ("retry.webm", b"retry", "audio/webm")},
        )
        new_after_end = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns/new-answer:transcribe",
            files={"file": ("late.webm", b"late", "audio/webm")},
        )
        assert finalized_after_end.status_code == after_end.status_code == 200
        finalized_after_end_turn = next(
            turn
            for turn in finalized_after_end.json()["turns"]
            if turn["client_turn_id"] == "pre-boundary-live"
        )
        after_end_turn = next(
            turn
            for turn in after_end.json()["turns"]
            if turn["client_turn_id"] == "candidate-answer"
        )
        assert finalized_after_end_turn["transcription_source"] == "final_model"
        assert after_end_turn == final_turn
        assert new_after_end.status_code == 409
        assert len(provider_calls) == 2

    with sqlite3.connect(database_path) as connection:
        events = connection.execute(
            "SELECT kind, quantity FROM usage_events WHERE session_id = ?",
            (interview["id"],),
        ).fetchall()
    assert events.count(("final_transcription_completed", 1)) == 2


def test_final_first_transcription_is_not_replaced_by_late_live_turn(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    transcription_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'final-first.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="server-key",
        azure_openai_realtime_deployment="interviewer-deployment",
        azure_openai_realtime_transcription_model="live-stt-deployment",
        azure_openai_final_transcription_deployment="final-stt-deployment",
    )

    async def fake_transcription(**_kwargs: object) -> FinalTranscription:
        return FinalTranscription(
            text="The final transcript arrived first.",
            deployment="final-stt-deployment",
            elapsed_ms=400,
            attempts=1,
        )

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    monkeypatch.setattr(
        "api.routes.realtime.transcribe_candidate_audio",
        fake_transcription,
    )
    with TestClient(create_app(transcription_settings)) as transcription_client:
        interview = _ready_session(transcription_client)
        _start_ready_interview(
            transcription_client, interview["id"], input_mode="voice"
        )

        def spooling_forbidden(*_args: object, **_kwargs: object) -> object:
            raise AssertionError("candidate audio must not use Starlette file spooling")

        monkeypatch.setattr(
            "starlette.formparsers.SpooledTemporaryFile", spooling_forbidden
        )
        finalized = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns/final-first:transcribe",
            data={
                "started_at": "2026-08-08T10:00:00Z",
                "ended_at": "2026-08-08T10:00:03Z",
            },
            files={"file": ("answer.ogg", b"audio", "audio/ogg")},
        )
        late_live = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json={
                "items": [
                    {
                        "client_turn_id": "final-first",
                        "speaker": "user",
                        "transcript": "The live transcript arrived late.",
                    }
                ]
            },
        )

    assert finalized.status_code == late_live.status_code == 200
    finalized_turn = finalized.json()["turns"][0]
    assert finalized_turn["started_at"] == "2026-08-08T10:00:00Z"
    assert finalized_turn["ended_at"] == "2026-08-08T10:00:03Z"
    assert len(late_live.json()["turns"]) == 1
    assert late_live.json()["turns"][0] == finalized_turn


def test_final_transcription_rejects_invalid_input_before_provider(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    transcription_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'transcription-input.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        enable_text_dev_mode=True,
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="server-key",
        azure_openai_realtime_deployment="interviewer-deployment",
        azure_openai_realtime_transcription_model="live-stt-deployment",
        azure_openai_final_transcription_deployment="final-stt-deployment",
        azure_openai_final_transcription_max_bytes=4,
    )
    provider_calls = 0

    async def forbidden_provider(**_kwargs: object) -> FinalTranscription:
        nonlocal provider_calls
        provider_calls += 1
        raise AssertionError("validation must happen before provider invocation")

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    monkeypatch.setattr(
        "api.routes.realtime.transcribe_candidate_audio",
        forbidden_provider,
    )
    with TestClient(create_app(transcription_settings)) as transcription_client:
        text_interview = _ready_session(transcription_client)
        _start_ready_interview(
            transcription_client, text_interview["id"], input_mode="text_dev"
        )
        text_mode = transcription_client.post(
            f"/api/interviews/{text_interview['id']}/turns/answer:transcribe",
            files={"file": ("answer.webm", b"data", "audio/webm")},
        )

        voice_interview = _ready_session(transcription_client)
        _start_ready_interview(
            transcription_client, voice_interview["id"], input_mode="voice"
        )
        invalid_mime = transcription_client.post(
            f"/api/interviews/{voice_interview['id']}/turns/answer:transcribe",
            files={"file": ("answer.wav", b"data", "audio/wav")},
        )
        empty = transcription_client.post(
            f"/api/interviews/{voice_interview['id']}/turns/answer:transcribe",
            files={"file": ("answer.mp4", b"", "audio/mp4")},
        )
        oversized = transcription_client.post(
            f"/api/interviews/{voice_interview['id']}/turns/answer:transcribe",
            files={"file": ("answer.ogg", b"12345", "audio/ogg")},
        )

    assert text_mode.status_code == 409
    assert invalid_mime.status_code == 415
    assert empty.status_code == 422
    assert oversized.status_code == 413
    assert provider_calls == 0


def test_accept_live_finalizes_only_owned_live_candidate_turns(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    database_path = tmp_path / "accept-live.db"
    transcription_settings = _clerk_settings(
        database_url=f"sqlite+aiosqlite:///{database_path}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        enable_text_dev_mode=True,
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="server-key",
        azure_openai_realtime_deployment="interviewer-deployment",
        azure_openai_realtime_transcription_model="live-stt-deployment",
    )
    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    owner = _principal_headers("turn-owner", "owner@example.test")
    stranger = _principal_headers("turn-stranger", "stranger@example.test")
    with TestClient(create_app(transcription_settings)) as transcription_client:
        voice_interview = _ready_session(transcription_client, headers=owner)
        _start_ready_interview(
            transcription_client,
            voice_interview["id"],
            input_mode="voice",
            headers=owner,
        )
        turns = transcription_client.post(
            f"/api/interviews/{voice_interview['id']}/turns:batch",
            headers=owner,
            json={
                "items": [
                    {
                        "client_turn_id": "assistant-turn",
                        "speaker": "assistant",
                        "transcript": "Tell me about a project.",
                    },
                    {
                        "client_turn_id": "live-turn",
                        "speaker": "user",
                        "transcript": "I built a reliable API.",
                    },
                ]
            },
        )
        assert turns.status_code == 200
        accepted = transcription_client.post(
            f"/api/interviews/{voice_interview['id']}/turns/live-turn:accept-live",
            headers=owner,
        )
        repeated = transcription_client.post(
            f"/api/interviews/{voice_interview['id']}/turns/live-turn:accept-live",
            headers=owner,
        )
        assistant = transcription_client.post(
            f"/api/interviews/{voice_interview['id']}/turns/assistant-turn:accept-live",
            headers=owner,
        )
        missing = transcription_client.post(
            f"/api/interviews/{voice_interview['id']}/turns/missing-turn:accept-live",
            headers=owner,
        )
        foreign = transcription_client.post(
            f"/api/interviews/{voice_interview['id']}/turns/live-turn:accept-live",
            headers=stranger,
        )

        text_interview = _ready_session(transcription_client, headers=owner)
        _start_ready_interview(
            transcription_client,
            text_interview["id"],
            input_mode="text_dev",
            headers=owner,
        )
        transcription_client.post(
            f"/api/interviews/{text_interview['id']}/turns:batch",
            headers=owner,
            json={
                "items": [
                    {
                        "client_turn_id": "typed-turn",
                        "speaker": "user",
                        "transcript": "This was typed.",
                    }
                ]
            },
        )
        typed = transcription_client.post(
            f"/api/interviews/{text_interview['id']}/turns/typed-turn:accept-live",
            headers=owner,
        )

    assert accepted.status_code == repeated.status_code == 200
    accepted_turn = next(
        turn
        for turn in accepted.json()["turns"]
        if turn["client_turn_id"] == "live-turn"
    )
    repeated_turn = next(
        turn
        for turn in repeated.json()["turns"]
        if turn["client_turn_id"] == "live-turn"
    )
    assert accepted_turn == repeated_turn
    assert accepted_turn["transcription_source"] == "realtime_live"
    assert accepted_turn["transcript"] == "I built a reliable API."
    assert accepted_turn["transcription_finalized_at"] is not None
    assert assistant.status_code == typed.status_code == 409
    assert missing.status_code == foreign.status_code == 404
    with sqlite3.connect(database_path) as connection:
        fallback_events = connection.execute(
            "SELECT kind, quantity FROM usage_events WHERE session_id = ?",
            (voice_interview["id"],),
        ).fetchall()
    assert fallback_events.count(("live_transcription_fallback", 1)) == 1


def test_accept_live_returns_existing_final_runtime_without_fallback_event(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    database_path = tmp_path / "accept-existing-final.db"
    settings = _dual_transcription_settings(database_path)

    async def fake_transcription(**_kwargs: object) -> FinalTranscription:
        return FinalTranscription(
            text="The committed final transcript must remain authoritative.",
            deployment="final-stt-deployment",
            elapsed_ms=300,
            attempts=1,
        )

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    monkeypatch.setattr(
        "api.routes.realtime.transcribe_candidate_audio", fake_transcription
    )
    with TestClient(create_app(settings)) as transcription_client:
        interview = _ready_session(transcription_client)
        _start_ready_interview(
            transcription_client, interview["id"], input_mode="voice"
        )
        live = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json={
                "items": [
                    {
                        "client_turn_id": "lost-final-response",
                        "speaker": "user",
                        "transcript": "The live transcript arrived first.",
                    }
                ]
            },
        )
        assert live.status_code == 200
        finalized = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns/lost-final-response:transcribe",
            files={"file": ("answer.webm", b"audio", "audio/webm")},
        )
        accepted = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns/lost-final-response:accept-live"
        )

    assert finalized.status_code == accepted.status_code == 200
    finalized_turn = next(
        turn
        for turn in finalized.json()["turns"]
        if turn["client_turn_id"] == "lost-final-response"
    )
    accepted_turn = next(
        turn
        for turn in accepted.json()["turns"]
        if turn["client_turn_id"] == "lost-final-response"
    )
    assert accepted_turn == finalized_turn
    assert accepted_turn["transcript"] == (
        "The committed final transcript must remain authoritative."
    )
    assert accepted_turn["transcription_source"] == "final_model"
    with sqlite3.connect(database_path) as connection:
        fallback_events = connection.execute(
            "SELECT COUNT(*) FROM usage_events "
            "WHERE session_id = ? AND kind = 'live_transcription_fallback'",
            (interview["id"],),
        ).fetchone()[0]
    assert fallback_events == 0


def test_concurrent_final_first_requests_converge_on_sqlite(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    database_path = tmp_path / "final-final-race.db"
    settings = _dual_transcription_settings(database_path)
    provider_barrier = threading.Barrier(2)
    second_read_barrier = threading.Barrier(2)
    provider_calls = 0

    async def synchronized_transcription(**_kwargs: object) -> FinalTranscription:
        nonlocal provider_calls
        provider_calls += 1
        await asyncio.to_thread(provider_barrier.wait, 10)
        return FinalTranscription(
            text="Both final requests converge here.",
            deployment="final-stt-deployment",
            elapsed_ms=10,
            attempts=1,
        )

    original_turn_lookup = realtime_routes._turn_by_client_id

    async def synchronized_turn_lookup(*args: object, **kwargs: object):
        turn = await original_turn_lookup(*args, **kwargs)
        if args[2] == "race-final" and provider_calls == 2 and turn is None:
            await asyncio.to_thread(second_read_barrier.wait, 10)
        return turn

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    monkeypatch.setattr(
        "api.routes.realtime.transcribe_candidate_audio", synchronized_transcription
    )
    monkeypatch.setattr(
        "api.routes.realtime._turn_by_client_id", synchronized_turn_lookup
    )
    with TestClient(
        create_app(settings), raise_server_exceptions=False
    ) as transcription_client:
        interview = _ready_session(transcription_client)
        _start_ready_interview(
            transcription_client, interview["id"], input_mode="voice"
        )

        def transcribe() -> object:
            return transcription_client.post(
                f"/api/interviews/{interview['id']}/turns/race-final:transcribe",
                files={"file": ("answer.webm", b"audio", "audio/webm")},
            )

        with ThreadPoolExecutor(max_workers=2) as executor:
            responses = [
                future.result()
                for future in [executor.submit(transcribe) for _ in range(2)]
            ]

    assert [response.status_code for response in responses] == [200, 200]
    response_turns = [
        next(
            turn
            for turn in response.json()["turns"]
            if turn["client_turn_id"] == "race-final"
        )
        for response in responses
    ]
    assert response_turns[0]["id"] == response_turns[1]["id"]
    with sqlite3.connect(database_path) as connection:
        turn_count = connection.execute(
            "SELECT COUNT(*) FROM interview_turns WHERE client_turn_id = 'race-final'"
        ).fetchone()[0]
        event_count = connection.execute(
            "SELECT COUNT(*) FROM usage_events "
            "WHERE kind = 'final_transcription_completed'"
        ).fetchone()[0]
    assert turn_count == event_count == 1


def test_concurrent_live_and_final_inserts_converge_on_sqlite(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    database_path = tmp_path / "live-final-race.db"
    settings = _dual_transcription_settings(database_path)
    provider_started = threading.Event()
    release_provider = threading.Event()
    batch_waiting = threading.Event()
    release_batch = threading.Event()

    async def synchronized_transcription(**_kwargs: object) -> FinalTranscription:
        provider_started.set()
        assert await asyncio.to_thread(release_provider.wait, 10)
        return FinalTranscription(
            text="The final transcript wins the race.",
            deployment="final-stt-deployment",
            elapsed_ms=10,
            attempts=1,
        )

    original_turn_lookup = realtime_routes._turn_by_client_id

    async def release_batch_after_final_read(*args: object, **kwargs: object):
        turn = await original_turn_lookup(*args, **kwargs)
        if args[2] == "race-live-final" and release_provider.is_set() and turn is None:
            release_batch.set()
        return turn

    original_commit = AsyncSession.commit

    async def hold_live_insert_commit(session: AsyncSession) -> None:
        if any(
            isinstance(item, InterviewTurn)
            and item.client_turn_id == "race-live-final"
            and item.transcription_source == "realtime_live"
            for item in session.new
        ):
            batch_waiting.set()
            assert await asyncio.to_thread(release_batch.wait, 10)
        await original_commit(session)

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    monkeypatch.setattr(
        "api.routes.realtime.transcribe_candidate_audio", synchronized_transcription
    )
    monkeypatch.setattr(
        "api.routes.realtime._turn_by_client_id", release_batch_after_final_read
    )
    with TestClient(
        create_app(settings), raise_server_exceptions=False
    ) as transcription_client:
        interview = _ready_session(transcription_client)
        _start_ready_interview(
            transcription_client, interview["id"], input_mode="voice"
        )
        monkeypatch.setattr(AsyncSession, "commit", hold_live_insert_commit)

        def transcribe() -> object:
            return transcription_client.post(
                f"/api/interviews/{interview['id']}/turns/race-live-final:transcribe",
                files={"file": ("answer.webm", b"audio", "audio/webm")},
            )

        def save_live() -> object:
            return transcription_client.post(
                f"/api/interviews/{interview['id']}/turns:batch",
                json={
                    "items": [
                        {
                            "client_turn_id": "race-live-final",
                            "speaker": "user",
                            "transcript": "The live transcript loses the race.",
                        }
                    ]
                },
            )

        with ThreadPoolExecutor(max_workers=2) as executor:
            final_future = executor.submit(transcribe)
            assert provider_started.wait(10)
            live_future = executor.submit(save_live)
            assert batch_waiting.wait(10)
            release_provider.set()
            responses = [final_future.result(), live_future.result()]

    assert [response.status_code for response in responses] == [200, 200]
    with sqlite3.connect(database_path) as connection:
        saved = connection.execute(
            "SELECT transcript, transcription_source FROM interview_turns "
            "WHERE client_turn_id = 'race-live-final'"
        ).fetchall()
    assert saved == [("The final transcript wins the race.", "final_model")]


def test_concurrent_live_acceptance_writes_one_fallback_event_on_sqlite(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    database_path = tmp_path / "accept-accept-race.db"
    settings = _dual_transcription_settings(database_path)
    read_barrier = threading.Barrier(2)
    original_turn_lookup = realtime_routes._turn_by_client_id

    async def synchronized_turn_lookup(*args: object, **kwargs: object):
        turn = await original_turn_lookup(*args, **kwargs)
        if (
            args[2] == "race-accept"
            and turn is not None
            and turn.transcription_finalized_at is None
        ):
            await asyncio.to_thread(read_barrier.wait, 10)
        return turn

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    monkeypatch.setattr(
        "api.routes.realtime._turn_by_client_id", synchronized_turn_lookup
    )
    with TestClient(
        create_app(settings), raise_server_exceptions=False
    ) as transcription_client:
        interview = _ready_session(transcription_client)
        _start_ready_interview(
            transcription_client, interview["id"], input_mode="voice"
        )
        saved = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json={
                "items": [
                    {
                        "client_turn_id": "race-accept",
                        "speaker": "user",
                        "transcript": "Accept this live transcript once.",
                    }
                ]
            },
        )
        assert saved.status_code == 200

        def accept() -> object:
            return transcription_client.post(
                f"/api/interviews/{interview['id']}/turns/race-accept:accept-live"
            )

        with ThreadPoolExecutor(max_workers=2) as executor:
            responses = [
                future.result()
                for future in [executor.submit(accept) for _ in range(2)]
            ]

    assert [response.status_code for response in responses] == [200, 200]
    with sqlite3.connect(database_path) as connection:
        event_count = connection.execute(
            "SELECT COUNT(*) FROM usage_events "
            "WHERE kind = 'live_transcription_fallback'"
        ).fetchone()[0]
    assert event_count == 1


def test_transcription_errors_and_telemetry_are_content_free(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, caplog: pytest.LogCaptureFixture
) -> None:
    database_path = tmp_path / "transcription-telemetry.db"
    transcription_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{database_path}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="server-key",
        azure_openai_realtime_deployment="interviewer-deployment",
        azure_openai_realtime_transcription_model="live-stt-deployment",
        azure_openai_final_transcription_deployment="final-stt-deployment",
    )
    provider_attempt = 0

    async def failed_transcription(**_kwargs: object) -> FinalTranscription:
        nonlocal provider_attempt
        provider_attempt += 1
        if provider_attempt == 2:
            raise RuntimeError("provider response body: secret upstream payload")
        raise TranscriptionServiceError(
            "Final transcription is temporarily unavailable. Try again.",
            code="transcription_unavailable",
            status_code=502,
            attempts=3,
        )

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    monkeypatch.setattr(
        "api.routes.realtime.transcribe_candidate_audio",
        failed_transcription,
    )
    caplog.set_level("INFO")
    with TestClient(create_app(transcription_settings)) as transcription_client:
        interview = _ready_session(transcription_client)
        _start_ready_interview(
            transcription_client, interview["id"], input_mode="voice"
        )
        failure = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns/failed-turn:transcribe",
            files={
                "file": (
                    "private-answer.webm",
                    b"private-candidate-audio",
                    "audio/webm",
                )
            },
        )
        unexpected_failure = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns/unexpected-failure:transcribe",
            files={"file": ("answer.ogg", b"other-private-audio", "audio/ogg")},
        )
        allowed = transcription_client.post(
            f"/api/interviews/{interview['id']}/transcription-events",
            json={"kind": "live_transcription_completed"},
        )
        invalid = transcription_client.post(
            f"/api/interviews/{interview['id']}/transcription-events",
            json={"kind": "arbitrary_event"},
        )
        transcript_field = transcription_client.post(
            f"/api/interviews/{interview['id']}/transcription-events",
            json={
                "kind": "live_transcription_failed",
                "transcript": "private candidate transcript",
            },
        )
        audio_field = transcription_client.post(
            f"/api/interviews/{interview['id']}/transcription-events",
            json={
                "kind": "double_transcription_failure",
                "audio": "private-candidate-audio",
            },
        )

    assert failure.status_code == 502
    assert failure.headers["X-Error-ID"] == failure.json()["error"]["id"]
    assert failure.json()["error"]["code"] == "http_error"
    assert "provider response body" not in failure.text
    assert unexpected_failure.status_code == 502
    assert (
        unexpected_failure.headers["X-Error-ID"]
        == unexpected_failure.json()["error"]["id"]
    )
    assert "provider response body" not in unexpected_failure.text
    assert allowed.status_code == 204
    assert invalid.status_code == transcript_field.status_code == 422
    assert audio_field.status_code == 422
    assert "private-candidate-audio" not in caplog.text
    assert "private candidate transcript" not in caplog.text
    assert "provider response body" not in caplog.text
    with sqlite3.connect(database_path) as connection:
        events = connection.execute(
            "SELECT kind, quantity FROM usage_events WHERE session_id = ?",
            (interview["id"],),
        ).fetchall()
    assert events.count(("live_transcription_completed", 1)) == 1
    assert not {
        "arbitrary_event",
        "live_transcription_failed",
        "double_transcription_failure",
    } & {kind for kind, _quantity in events}


def test_voice_evaluation_waits_for_transcription_finalization(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    transcription_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'pending-finalization.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="server-key",
        azure_openai_realtime_deployment="interviewer-deployment",
        azure_openai_realtime_transcription_model="live-stt-deployment",
        azure_openai_text_deployment="evaluation-deployment",
    )
    evaluation_calls = 0

    async def forbidden_evaluation(**_kwargs: object) -> EvaluationReport:
        nonlocal evaluation_calls
        evaluation_calls += 1
        raise AssertionError("pending transcription must stop before evaluation")

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    monkeypatch.setattr(
        "api.services.evaluation_jobs.evaluate_transcript", forbidden_evaluation
    )
    with TestClient(create_app(transcription_settings)) as transcription_client:
        interview = _ready_session(transcription_client)
        _start_ready_interview(
            transcription_client, interview["id"], input_mode="voice"
        )
        saved = transcription_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json={
                "items": [
                    {
                        "client_turn_id": "pending-final",
                        "speaker": "user",
                        "transcript": "The live transcript is acknowledged.",
                        "delivery_status": "acknowledged",
                    }
                ]
            },
        )
        completed = transcription_client.post(
            f"/api/interviews/{interview['id']}/complete"
        )
        runtime = transcription_client.get(f"/api/interviews/{interview['id']}/runtime")

    assert saved.status_code == completed.status_code == 200
    assert saved.json()["turns"][0]["delivery_status"] == "acknowledged"
    assert saved.json()["turns"][0]["transcription_finalized_at"] is None
    assert runtime.json()["status"] == "FAILED_RECOVERABLE"
    assert evaluation_calls == 0


def test_assistant_only_transcript_never_reaches_evaluator(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    settings = _dual_transcription_settings(tmp_path / "assistant-only.db")
    evaluation_calls = 0

    async def forbidden_evaluation(**_kwargs: object) -> EvaluationReport:
        nonlocal evaluation_calls
        evaluation_calls += 1
        raise AssertionError("assistant-only transcript must not reach evaluation")

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    monkeypatch.setattr(
        "api.services.evaluation_jobs.evaluate_transcript", forbidden_evaluation
    )
    with TestClient(create_app(settings)) as evaluation_client:
        interview = _ready_session(evaluation_client)
        _start_ready_interview(evaluation_client, interview["id"], input_mode="voice")
        saved = evaluation_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json={
                "items": [
                    {
                        "client_turn_id": "assistant-only",
                        "speaker": "assistant",
                        "transcript": "Tell me about an API you designed.",
                    }
                ]
            },
        )
        completed = evaluation_client.post(
            f"/api/interviews/{interview['id']}/complete"
        )
        runtime = evaluation_client.get(f"/api/interviews/{interview['id']}/runtime")

    assert saved.status_code == completed.status_code == 200
    assert runtime.json()["status"] == "FAILED_RECOVERABLE"
    assert evaluation_calls == 0


def test_legacy_voice_candidate_turn_remains_evaluable(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    database_path = tmp_path / "legacy-voice-evaluation.db"
    settings = _dual_transcription_settings(database_path)
    evaluation_calls = 0

    async def observed_evaluation(**_kwargs: object) -> EvaluationReport:
        nonlocal evaluation_calls
        evaluation_calls += 1
        raise AssertionError("legacy compatibility reached the evaluator")

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", _fake_realtime_secret
    )
    monkeypatch.setattr(
        "api.services.evaluation_jobs.evaluate_transcript", observed_evaluation
    )
    with TestClient(create_app(settings)) as evaluation_client:
        interview = _ready_session(evaluation_client)
        _start_ready_interview(evaluation_client, interview["id"], input_mode="voice")
        saved = evaluation_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json={
                "items": [
                    {
                        "client_turn_id": "legacy-candidate",
                        "speaker": "user",
                        "transcript": "This predates dual transcription.",
                    }
                ]
            },
        )
        assert saved.status_code == 200
        with sqlite3.connect(database_path) as connection:
            connection.execute(
                "UPDATE interview_turns SET transcription_source = 'legacy', "
                "transcription_model = NULL, transcription_finalized_at = NULL "
                "WHERE session_id = ? AND client_turn_id = 'legacy-candidate'",
                (interview["id"],),
            )
            connection.commit()
        completed = evaluation_client.post(
            f"/api/interviews/{interview['id']}/complete"
        )

    assert completed.status_code == 200
    assert evaluation_calls == 1


def test_m3_text_realtime_flow_is_private_long_and_idempotent(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    realtime_settings = Settings(
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'm3.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        enable_text_dev_mode=True,
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="permanent-server-key",
        azure_openai_realtime_deployment="realtime-deployment",
    )
    captured: dict[str, object] = {}

    async def fake_secret(**kwargs):
        captured.update(kwargs)
        return RealtimeClientSecret(
            value="ek_temporary",
            expires_at=1_786_000_000,
            calls_url=(
                "https://example.services.ai.azure.com/openai/v1/"
                "realtime/calls?webrtcfilter=on"
            ),
        )

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", fake_secret
    )
    with TestClient(create_app(realtime_settings)) as realtime_client:
        interview = _ready_session(realtime_client)
        capabilities = realtime_client.get("/api/capabilities").json()
        assert capabilities == {
            "text_dev_mode_enabled": True,
            "realtime_configured": True,
            "live_transcription_configured": False,
            "final_transcription_configured": False,
            "typed_answer_max_characters": 20_000,
            "supported_durations": [15, 30, 45, 60],
        }

        secret = realtime_client.post(
            f"/api/interviews/{interview['id']}/realtime-client-secret",
            json={
                "input_mode": "text_dev",
                "duration_minutes": 15,
                "interview_type": "technical_behavioral",
            },
        )
        assert secret.status_code == 200
        assert set(secret.json()) == {
            "client_secret",
            "expires_at",
            "calls_url",
            "input_mode",
            "prompt_version",
        }
        assert "permanent-server-key" not in secret.text
        assert "server-owned" not in secret.text
        assert captured["input_mode"] == "text_dev"
        assert "TRUSTED_SESSION_CONTEXT_JSON" in captured["instructions"]

        connected = realtime_client.post(
            f"/api/interviews/{interview['id']}/connection-state",
            json={"state": "connected"},
        )
        assert connected.status_code == 200
        assert connected.json()["status"] == "IN_PROGRESS"
        assert connected.json()["started_at"]
        assert connected.json()["ends_at"]

        long_answer = "🙂" * 20_000
        pending_payload = {
            "items": [
                {
                    "client_turn_id": "item_long_answer",
                    "speaker": "user",
                    "transcript": long_answer,
                    "delivery_status": "pending",
                }
            ]
        }
        pending = realtime_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json=pending_payload,
        )
        assert pending.status_code == 200
        assert len(pending.json()["turns"]) == 1
        legacy_turn = pending.json()["turns"][0]
        assert legacy_turn["delivery_status"] == "pending"
        assert legacy_turn["transcription_source"] == "typed"
        assert legacy_turn["transcription_model"] is None
        assert legacy_turn["transcription_finalized_at"] is not None

        pending_payload["items"][0]["delivery_status"] = "acknowledged"
        acknowledged = realtime_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json=pending_payload,
        )
        repeated = realtime_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json=pending_payload,
        )
        assert acknowledged.status_code == repeated.status_code == 200
        assert len(repeated.json()["turns"]) == 1
        assert repeated.json()["turns"][0]["delivery_status"] == "acknowledged"

        conflicting = pending_payload.copy()
        conflicting["items"] = [
            {**pending_payload["items"][0], "transcript": "different"}
        ]
        assert (
            realtime_client.post(
                f"/api/interviews/{interview['id']}/turns:batch",
                json=conflicting,
            ).status_code
            == 409
        )

        completed = realtime_client.post(f"/api/interviews/{interview['id']}/complete")
        completed_again = realtime_client.post(
            f"/api/interviews/{interview['id']}/complete"
        )
        assert completed.json()["status"] == "TRANSCRIPT_FINALIZING"
        assert completed_again.json()["started_at"] == completed.json()["started_at"]


def test_m3_reconnect_rejects_an_expired_recovery_window(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    realtime_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'm3-reconnect.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        enable_text_dev_mode=True,
        realtime_reconnect_window_seconds=-1,
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="permanent-server-key",
        azure_openai_realtime_deployment="realtime-deployment",
    )

    async def fake_secret(**_kwargs):
        return RealtimeClientSecret(
            value="ek_temporary",
            expires_at=1_786_000_000,
            calls_url=(
                "https://example.services.ai.azure.com/openai/v1/"
                "realtime/calls?webrtcfilter=on"
            ),
        )

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", fake_secret
    )
    with TestClient(create_app(realtime_settings)) as realtime_client:
        interview = _ready_session(realtime_client)
        request_payload = {
            "input_mode": "text_dev",
            "duration_minutes": 15,
            "interview_type": "technical_behavioral",
        }
        assert (
            realtime_client.post(
                f"/api/interviews/{interview['id']}/realtime-client-secret",
                json=request_payload,
            ).status_code
            == 200
        )
        realtime_client.post(
            f"/api/interviews/{interview['id']}/connection-state",
            json={"state": "connected"},
        )
        realtime_client.post(
            f"/api/interviews/{interview['id']}/connection-state",
            json={"state": "reconnecting"},
        )

        expired = realtime_client.post(
            f"/api/interviews/{interview['id']}/realtime-client-secret",
            json=request_payload,
        )
        assert expired.status_code == 409
        assert "recovery window has expired" in expired.json()["error"]["message"]


def test_m3_rejects_invalid_state_transitions(client: TestClient) -> None:
    interview = _create_session(client)

    connected = client.post(
        f"/api/interviews/{interview['id']}/connection-state",
        json={"state": "connected"},
    )
    completed = client.post(f"/api/interviews/{interview['id']}/complete")

    assert connected.status_code == 409
    assert completed.status_code == 409


def test_m3_timer_is_server_authoritative_and_utc_on_refresh(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    database_path = tmp_path / "m3-timer.db"
    realtime_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{database_path}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        enable_text_dev_mode=True,
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="server-key",
        azure_openai_realtime_deployment="realtime-deployment",
    )

    async def fake_secret(**_kwargs):
        return RealtimeClientSecret(
            value="ek_temporary",
            expires_at=2_000_000_000,
            calls_url="https://example.services.ai.azure.com/openai/v1/realtime/calls",
        )

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", fake_secret
    )
    with TestClient(create_app(realtime_settings)) as realtime_client:
        interview = _ready_session(realtime_client)
        secret_payload = {
            "input_mode": "text_dev",
            "duration_minutes": 15,
            "interview_type": "technical_behavioral",
        }
        assert (
            realtime_client.post(
                f"/api/interviews/{interview['id']}/realtime-client-secret",
                json=secret_payload,
            ).status_code
            == 200
        )
        connected = realtime_client.post(
            f"/api/interviews/{interview['id']}/connection-state",
            json={"state": "connected"},
        )
        assert connected.json()["started_at"].endswith("Z")

        refreshed = realtime_client.get(f"/api/interviews/{interview['id']}/runtime")
        assert refreshed.json()["started_at"].endswith("Z")
        assert refreshed.json()["ends_at"].endswith("Z")

        expired_start = datetime.now(UTC) - timedelta(minutes=16)
        with sqlite3.connect(database_path) as connection:
            connection.execute(
                "UPDATE interview_sessions SET started_at = ? WHERE id = ?",
                (expired_start.replace(tzinfo=None), interview["id"]),
            )
            connection.commit()

        expired_runtime = realtime_client.get(
            f"/api/interviews/{interview['id']}/runtime"
        )
        assert expired_runtime.json()["status"] == "TRANSCRIPT_FINALIZING"
        assert (
            realtime_client.post(
                f"/api/interviews/{interview['id']}/realtime-client-secret",
                json=secret_payload,
            ).status_code
            == 409
        )
        assert (
            realtime_client.post(
                f"/api/interviews/{interview['id']}/turns:batch",
                json={
                    "items": [
                        {
                            "client_turn_id": "late_turn",
                            "speaker": "user",
                            "transcript": "This arrived after the server deadline.",
                        }
                    ]
                },
            ).status_code
            == 409
        )


def test_m3_start_is_idempotent_and_setup_freezes(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    realtime_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'm3-freeze.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        enable_text_dev_mode=True,
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="server-key",
        azure_openai_realtime_deployment="realtime-deployment",
    )
    calls = 0

    async def fake_secret(**_kwargs):
        nonlocal calls
        calls += 1
        return RealtimeClientSecret(
            value="ek_same_attempt",
            expires_at=2_000_000_000,
            calls_url="https://example.services.ai.azure.com/openai/v1/realtime/calls",
        )

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", fake_secret
    )
    with TestClient(create_app(realtime_settings)) as realtime_client:
        interview = _ready_session(realtime_client)
        setup = realtime_client.get(f"/api/interviews/{interview['id']}/setup").json()
        secret_payload = {
            "input_mode": "text_dev",
            "duration_minutes": 15,
            "interview_type": "technical_behavioral",
        }
        first = realtime_client.post(
            f"/api/interviews/{interview['id']}/realtime-client-secret",
            json=secret_payload,
        )
        repeated = realtime_client.post(
            f"/api/interviews/{interview['id']}/realtime-client-secret",
            json=secret_payload,
        )
        assert first.json() == repeated.json()
        assert calls == 1

        profile = setup["profile"]
        profile_edit = {
            "headline": profile["headline"],
            "claims": [
                {"id": claim["id"], "text": claim["text"]}
                for claim in profile["claims"]
            ],
        }
        scorecard = setup["scorecard"]
        scorecard_edit = {
            "competencies": [
                {
                    key: competency[key]
                    for key in (
                        "id",
                        "name",
                        "description",
                        "weight",
                        "classification",
                        "seniority_expectation",
                        "evidence_to_collect",
                        "question_families",
                    )
                }
                for competency in scorecard["competencies"]
            ]
        }
        assert (
            realtime_client.patch(
                f"/api/candidate-profiles/{profile['id']}", json=profile_edit
            ).status_code
            == 409
        )
        assert (
            realtime_client.patch(
                f"/api/scorecards/{scorecard['id']}", json=scorecard_edit
            ).status_code
            == 409
        )


def test_database_engine_hides_private_parameters(client: TestClient) -> None:
    assert client.app.state.engine.sync_engine.hide_parameters is True


def test_clerk_session_token_creates_the_signed_in_user(tmp_path: Path) -> None:
    clerk_settings = _clerk_settings(
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'clerk.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
    )
    headers = _principal_headers("user_external", "candidate@example.test")

    with TestClient(create_app(clerk_settings)) as clerk_client:
        response = clerk_client.get("/api/auth/me", headers=headers)

    assert response.status_code == 200
    assert response.json()["email"] == "candidate@example.test"
    assert response.json()["display_name"] == "candidate"


def test_managed_users_cannot_see_each_others_sessions(tmp_path: Path) -> None:
    clerk_settings = _clerk_settings(
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'ownership.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
    )
    with TestClient(create_app(clerk_settings)) as clerk_client:
        first_user = _principal_headers("user-a", "a@example.test")
        second_user = _principal_headers("user-b", "b@example.test")
        assert (
            clerk_client.post(
                "/api/interviews",
                headers=first_user,
                json={"title": "Private session"},
            ).status_code
            == 201
        )

        second_user_sessions = clerk_client.get("/api/interviews", headers=second_user)
        assert second_user_sessions.json() == {"items": []}


def test_unauthenticated_managed_request_has_error_id(tmp_path: Path) -> None:
    clerk_settings = _clerk_settings(
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'auth.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
    )
    with TestClient(create_app(clerk_settings)) as clerk_client:
        response = clerk_client.get("/api/auth/me")

    assert response.status_code == 401
    assert response.headers["X-Error-ID"] == response.json()["error"]["id"]


def test_clerk_ignores_a_session_cookie_without_a_bearer_header(
    tmp_path: Path,
) -> None:
    """A ``__session`` cookie must not be treated as a session token.

    Clerk's SDK sets that cookie on the origin, but its value is not reliably a
    JWT. Decoding it raised "Invalid header string" and surfaced as a spurious
    401 whenever a request beat the client to attaching its bearer header.
    """

    clerk_settings = _clerk_settings(
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'cookie.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
    )
    private_pem, _ = _clerk_keys()
    token = jwt.encode(
        _clerk_claims("user-cookie", "cookie@example.test"),
        private_pem,
        algorithm="RS256",
    )

    with TestClient(create_app(clerk_settings)) as clerk_client:
        # Even a *valid* token in the cookie is ignored; only the header counts.
        clerk_client.cookies.set("__session", token)
        assert clerk_client.get("/api/auth/me").status_code == 401

        # Non-JWT cookie content must not raise or leak a 500 either.
        clerk_client.cookies.set("__session", "not-a-jwt-at-all")
        assert clerk_client.get("/api/auth/me").status_code == 401

        # The header still authenticates with the cookie present.
        response = clerk_client.get(
            "/api/auth/me", headers={"Authorization": f"Bearer {token}"}
        )

    assert response.status_code == 200
    assert response.json()["email"] == "cookie@example.test"


def test_clerk_rejects_a_token_signed_by_another_key(tmp_path: Path) -> None:
    clerk_settings = _clerk_settings(
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'forged.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
    )
    impostor = rsa.generate_private_key(public_exponent=65537, key_size=2048)
    forged = jwt.encode(
        _clerk_claims("user-forged", "forged@example.test"),
        impostor.private_bytes(
            serialization.Encoding.PEM,
            serialization.PrivateFormat.PKCS8,
            serialization.NoEncryption(),
        ).decode(),
        algorithm="RS256",
    )

    with TestClient(create_app(clerk_settings)) as clerk_client:
        response = clerk_client.get(
            "/api/auth/me", headers={"Authorization": f"Bearer {forged}"}
        )

    assert response.status_code == 401


def test_clerk_rejects_a_token_from_an_unlisted_authorized_party(
    tmp_path: Path,
) -> None:
    clerk_settings = _clerk_settings(
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'azp.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
    )
    private_pem, _ = _clerk_keys()
    claims = _clerk_claims("user-elsewhere", "elsewhere@example.test")
    claims["azp"] = "https://someone-elses-app.test"
    token = jwt.encode(claims, private_pem, algorithm="RS256")

    with TestClient(create_app(clerk_settings)) as clerk_client:
        response = clerk_client.get(
            "/api/auth/me", headers={"Authorization": f"Bearer {token}"}
        )

    assert response.status_code == 401


def test_clerk_rejects_an_expired_token(tmp_path: Path) -> None:
    clerk_settings = _clerk_settings(
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'expired.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
    )
    private_pem, _ = _clerk_keys()
    claims = _clerk_claims("user-stale", "stale@example.test")
    claims["exp"] = int(time.time()) - 60
    token = jwt.encode(claims, private_pem, algorithm="RS256")

    with TestClient(create_app(clerk_settings)) as clerk_client:
        response = clerk_client.get(
            "/api/auth/me", headers={"Authorization": f"Bearer {token}"}
        )

    assert response.status_code == 401


def test_clerk_rejects_a_token_without_an_email_claim(tmp_path: Path) -> None:
    """Clerk's default session token omits email until the claim is added."""

    clerk_settings = _clerk_settings(
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'no-email.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
    )
    private_pem, _ = _clerk_keys()
    claims = _clerk_claims("user-anon", "anon@example.test")
    del claims["email"]
    token = jwt.encode(claims, private_pem, algorithm="RS256")

    with TestClient(create_app(clerk_settings)) as clerk_client:
        response = clerk_client.get(
            "/api/auth/me", headers={"Authorization": f"Bearer {token}"}
        )

    assert response.status_code == 401


def test_validation_errors_have_safe_error_ids(client: TestClient) -> None:
    response = client.post("/api/interviews", json={"title": "   "})

    assert response.status_code == 422
    assert response.headers["X-Error-ID"] == response.json()["error"]["id"]
    assert response.json()["error"]["message"] == (
        "The request did not pass validation."
    )


def test_react_build_is_served_from_fastapi(tmp_path: Path) -> None:
    dist_directory = REPO_ROOT / "web" / "dist"
    assert (dist_directory / "index.html").is_file(), "Run npm run build first."
    static_settings = Settings(
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'static.db'}",
        auto_create_schema=True,
        web_dist_dir=dist_directory,
    )

    with TestClient(create_app(static_settings)) as static_client:
        response = static_client.get("/practice/example")

    assert response.status_code == 200
    assert "<title>AI Interview Coach</title>" in response.text


def test_m4_completion_generates_one_evidence_backed_report(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    evaluation_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'm4.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        enable_text_dev_mode=True,
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="permanent-server-key",
        azure_openai_realtime_deployment="realtime-deployment",
        azure_openai_text_deployment="text-deployment",
    )
    model_calls = 0

    async def fake_secret(**_kwargs):
        return RealtimeClientSecret(
            value="ek_temporary",
            expires_at=1_786_000_000,
            calls_url="https://example.invalid/openai/v1/realtime/calls",
        )

    async def fake_evaluation(**kwargs):
        nonlocal model_calls
        model_calls += 1
        scorecard = kwargs["scorecard"]
        turns = kwargs["turns"]
        candidate_turn = next(turn for turn in turns if turn.speaker == "user")
        results = [
            CompetencyEvaluation(
                competency_id=competency.id,
                assessment="scored",
                score=4,
                rating_confidence="high",
                evidence=[
                    EvidenceCitation(
                        turn_id=candidate_turn.id,
                        quote="idempotency keys",
                    )
                ],
                evidence_summary="The candidate described a concrete safeguard.",
                gaps=[],
                recommendations=["Compare storage and retry trade-offs."],
            )
            for competency in scorecard.competencies
        ]
        return EvaluationReport(
            evaluator_version="test-evaluator-v1",
            competency_results=results,
            overall_score=4,
            assessed_weight=100,
            total_weight=100,
            coverage_percentage=100,
            strength_competency_ids=[results[0].competency_id],
            gap_competency_ids=[],
            practice_exercises=[],
            evidence_locations=[],
            validation_attempts=1,
        )

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", fake_secret
    )
    monkeypatch.setattr(
        "api.services.evaluation_jobs.evaluate_transcript", fake_evaluation
    )
    with TestClient(create_app(evaluation_settings)) as evaluation_client:
        interview = _ready_session(evaluation_client)
        secret_payload = {
            "input_mode": "text_dev",
            "duration_minutes": 15,
            "interview_type": "technical_behavioral",
        }
        assert (
            evaluation_client.post(
                f"/api/interviews/{interview['id']}/realtime-client-secret",
                json=secret_payload,
            ).status_code
            == 200
        )
        assert (
            evaluation_client.post(
                f"/api/interviews/{interview['id']}/connection-state",
                json={"state": "connected"},
            ).status_code
            == 200
        )
        turn = evaluation_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json={
                "items": [
                    {
                        "client_turn_id": "candidate-answer-1",
                        "speaker": "user",
                        "transcript": (
                            "I used idempotency keys and a unique database "
                            "constraint to make retries safe."
                        ),
                        "delivery_status": "acknowledged",
                    }
                ]
            },
        )
        assert turn.status_code == 200

        first = evaluation_client.post(f"/api/interviews/{interview['id']}/complete")
        repeated = evaluation_client.post(f"/api/interviews/{interview['id']}/complete")
        report = evaluation_client.get(f"/api/interviews/{interview['id']}/report")

    assert first.status_code == repeated.status_code == report.status_code == 200
    assert model_calls == 1
    assert report.json()["status"] == "REPORT_READY"
    assert report.json()["overall_score"] == 4
    assert report.json()["coverage_percentage"] == 100
    assert report.json()["delivery_coaching"]["status"] == "unavailable"
    assert report.json()["delivery_coaching"]["unavailable_reason"] == "text_input_mode"
    assert all(item["evidence"] for item in report.json()["competency_results"])
    assert all(
        item["evidence"][0]["quote"] == "idempotency keys"
        for item in report.json()["competency_results"]
    )


def test_m5_delivery_consent_metrics_disable_and_delete(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    delivery_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'm5.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        azure_openai_endpoint="https://example.services.ai.azure.com",
        azure_openai_api_key="permanent-server-key",
        azure_openai_realtime_deployment="realtime-deployment",
    )

    async def fake_secret(**_kwargs):
        return RealtimeClientSecret(
            value="ek_temporary",
            expires_at=1_786_000_000,
            calls_url="https://example.invalid/openai/v1/realtime/calls",
        )

    monkeypatch.setattr(
        "api.routes.realtime.create_realtime_client_secret", fake_secret
    )
    with TestClient(create_app(delivery_settings)) as delivery_client:
        interview = _ready_session(delivery_client)
        consent = delivery_client.post(
            f"/api/interviews/{interview['id']}/delivery-consent",
            json={"enabled": True, "consent_version": "delivery-v1"},
        )
        assert consent.status_code == 200
        assert consent.json()["status"] == "collecting"
        delivery_client.post(
            f"/api/interviews/{interview['id']}/realtime-client-secret",
            json={
                "input_mode": "voice",
                "duration_minutes": 15,
                "interview_type": "technical_behavioral",
            },
        )
        delivery_client.post(
            f"/api/interviews/{interview['id']}/connection-state",
            json={"state": "connected"},
        )
        turns = delivery_client.post(
            f"/api/interviews/{interview['id']}/turns:batch",
            json={
                "items": [
                    {
                        "client_turn_id": "assistant-question",
                        "speaker": "assistant",
                        "transcript": "Describe an API you designed.",
                        "delivery_status": "acknowledged",
                        "started_at": "2026-08-07T10:00:00Z",
                        "ended_at": "2026-08-07T10:00:03Z",
                    },
                    {
                        "client_turn_id": "candidate-answer",
                        "speaker": "user",
                        "transcript": (
                            "Um I designed a payment API with idempotency keys "
                            "and clear transaction boundaries."
                        ),
                        "delivery_status": "acknowledged",
                        "started_at": "2026-08-07T10:00:04Z",
                        "ended_at": "2026-08-07T10:00:10Z",
                    },
                ]
            },
        ).json()["turns"]
        candidate_turn = next(item for item in turns if item["speaker"] == "user")
        observed = delivery_client.post(
            f"/api/interviews/{interview['id']}/delivery-observations",
            json={
                "items": [
                    {
                        "turn_id": candidate_turn["id"],
                        "speech_segments": [
                            {
                                "started_at": "2026-08-07T10:00:04Z",
                                "ended_at": "2026-08-07T10:00:06Z",
                            },
                            {
                                "started_at": "2026-08-07T10:00:07Z",
                                "ended_at": "2026-08-07T10:00:10Z",
                            },
                        ],
                    }
                ]
            },
        )
        assert observed.status_code == 200
        assert observed.json()["status"] == "available"
        assert observed.json()["metrics"][0]["pause_count"] == 1
        assert observed.json()["metrics"][0]["filler_count"] == 1

        disabled = delivery_client.post(
            f"/api/interviews/{interview['id']}/delivery-consent",
            json={"enabled": False, "consent_version": "delivery-v1"},
        )
        assert disabled.json()["status"] == "disabled"
        assert disabled.json()["metrics"]
        deleted = delivery_client.delete(
            f"/api/interviews/{interview['id']}/delivery-metrics"
        )
        assert deleted.json()["status"] == "deleted"
        assert deleted.json()["metrics"] == []


def test_migrations_upgrade_and_roll_back(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    database_path = tmp_path / "migration.db"
    monkeypatch.setenv("DATABASE_URL", f"sqlite+aiosqlite:///{database_path}")
    get_settings.cache_clear()
    configuration = Config(SERVER_ROOT / "alembic.ini")

    command.upgrade(configuration, "head")
    with sqlite3.connect(database_path) as connection:
        tables = {
            row[0]
            for row in connection.execute(
                "SELECT name FROM sqlite_master WHERE type = 'table'"
            )
        }
    assert {
        "users",
        "interview_sessions",
        "uploads",
        "candidate_profiles",
        "job_targets",
        "scorecards",
        "interview_turns",
        "evaluations",
        "delivery_coaching",
        "usage_events",
        "deletion_receipts",
        "alembic_version",
    } <= tables
    with sqlite3.connect(database_path) as connection:
        interview_turn_columns = {
            row[1] for row in connection.execute("PRAGMA table_info(interview_turns)")
        }
    assert {
        "transcription_source",
        "transcription_model",
        "transcription_finalized_at",
    } <= interview_turn_columns

    command.downgrade(configuration, "base")
    with sqlite3.connect(database_path) as connection:
        tables = {
            row[0]
            for row in connection.execute(
                "SELECT name FROM sqlite_master WHERE type = 'table'"
            )
        }
    assert "users" not in tables
    assert "interview_sessions" not in tables
    assert "uploads" not in tables
    assert "candidate_profiles" not in tables
    assert "job_targets" not in tables
    assert "scorecards" not in tables
    assert "interview_turns" not in tables
    assert "evaluations" not in tables
    assert "delivery_coaching" not in tables
    assert "usage_events" not in tables
    assert "deletion_receipts" not in tables
    get_settings.cache_clear()


def test_m6_daily_session_quota_and_usage_summary(tmp_path: Path) -> None:
    quota_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'quota.db'}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
        daily_interview_quota=2,
    )

    with TestClient(create_app(quota_settings)) as quota_client:
        assert _create_session(quota_client)["id"]
        assert _create_session(quota_client)["id"]

        blocked = quota_client.post("/api/interviews", json={"title": "Over quota"})
        assert blocked.status_code == 429
        assert blocked.headers["Retry-After"]

        usage = quota_client.get("/api/operations/usage")
        assert usage.status_code == 200
        assert usage.json()["daily_interview_quota"] == 2
        assert usage.json()["daily_interviews_used"] == 2
        assert usage.json()["events"]["session_created"] == 2
        assert usage.json()["estimated_cost_usd"] == "0.000000"


def test_m6_session_deletion_is_idempotent_and_removes_private_data(
    tmp_path: Path,
) -> None:
    database_path = tmp_path / "delete-session.db"
    deletion_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{database_path}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
    )

    with TestClient(create_app(deletion_settings)) as deletion_client:
        interview = _ready_session(deletion_client)
        response = deletion_client.request(
            "DELETE",
            f"/api/interviews/{interview['id']}",
            json={"confirmation": "DELETE"},
        )
        repeated = deletion_client.request(
            "DELETE",
            f"/api/interviews/{interview['id']}",
            json={"confirmation": "DELETE"},
        )

        assert response.status_code == repeated.status_code == 204
        assert deletion_client.get("/api/interviews").json() == {"items": []}

    with sqlite3.connect(database_path) as connection:
        for table in (
            "interview_sessions",
            "interview_turns",
            "evaluations",
            "delivery_coaching",
            "candidate_profiles",
            "uploads",
            "scorecards",
            "job_targets",
        ):
            assert (
                connection.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0] == 0
            )
        receipt = connection.execute(
            "SELECT kind, status FROM deletion_receipts"
        ).fetchone()
        assert receipt == ("session", "completed")


def test_m6_account_deletion_is_idempotent_and_blocks_recreation(
    tmp_path: Path,
) -> None:
    database_path = tmp_path / "delete-account.db"
    deletion_settings = Settings(
        _env_file=None,
        app_env="test",
        auth_mode="local",
        database_url=f"sqlite+aiosqlite:///{database_path}",
        auto_create_schema=True,
        web_dist_dir=tmp_path / "missing-dist",
    )

    with TestClient(create_app(deletion_settings)) as deletion_client:
        _ready_session(deletion_client)
        response = deletion_client.request(
            "DELETE",
            "/api/account",
            json={"confirmation": "DELETE MY ACCOUNT"},
        )
        repeated = deletion_client.request(
            "DELETE",
            "/api/account",
            json={"confirmation": "DELETE MY ACCOUNT"},
        )

        assert response.status_code == repeated.status_code == 204
        assert deletion_client.get("/api/auth/me").status_code == 410

    with sqlite3.connect(database_path) as connection:
        assert connection.execute("SELECT COUNT(*) FROM users").fetchone()[0] == 0
        receipt = connection.execute(
            "SELECT kind, status FROM deletion_receipts WHERE kind = 'account'"
        ).fetchone()
        assert receipt == ("account", "completed")
